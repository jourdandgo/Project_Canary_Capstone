"""Build the Project Canary defense cheat sheet DOCX from its Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "PROJECT_CANARY_DEFENSE_CHEAT_SHEET.md"
OUTPUT = ROOT / "docs" / "Project_Canary_Defense_Cheat_Sheet.docx"

GREEN = RGBColor(23, 63, 49)
GREEN_2 = RGBColor(40, 98, 69)
LIME = RGBColor(150, 190, 43)
INK = RGBColor(42, 55, 48)
MUTED = RGBColor(96, 112, 105)
WHITE = RGBColor(255, 255, 255)
LIGHT = "EDF5F0"
LINE = "D7E4DC"


def set_font(run, size: float = 11, color: RGBColor = INK, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table_pr = table._tbl.tblPr
    table_w = table_pr.find(qn("w:tblW"))
    if table_w is None:
        table_w = OxmlElement("w:tblW")
        table_pr.append(table_w)
    table_w.set(qn("w:w"), str(sum(widths)))
    table_w.set(qn("w:type"), "dxa")
    table_ind = table_pr.find(qn("w:tblInd"))
    if table_ind is None:
        table_ind = OxmlElement("w:tblInd")
        table_pr.append(table_ind)
    table_ind.set(qn("w:w"), "120")
    table_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_inline(paragraph, text: str, *, size: float = 11, color: RGBColor = INK) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=size, color=color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=size - 0.5, color=GREEN_2, bold=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size, color=color)


def paragraph_spacing(paragraph, before: float = 0, after: float = 6, line: float = 1.25) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Project Canary · Defense Cheat Sheet  |  ")
    set_font(run, size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def new_numbering_sequence(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (("w:start", "1"), ("w:numFmt", "decimal"), ("w:lvlText", "%1."), ("w:lvlJc", "left"), ("w:suff", "space")):
        node = OxmlElement(tag)
        node.set(qn("w:val"), value)
        level.append(node)
    p_pr = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend([level, number])
    p_pr.append(num_pr)
    paragraph_spacing(paragraph, after=3, line=1.18)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    heading_tokens = {
        "Heading 1": (16, GREEN, 16, 7),
        "Heading 2": (13, GREEN_2, 12, 5),
        "Heading 3": (11.5, GREEN, 9, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.18

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("PROJECT CANARY  ·  CAPSTONE DEFENSE REFERENCE")
    set_font(hr, size=8.5, color=GREEN, bold=True)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    paragraph_spacing(kicker, after=2)
    kr = kicker.add_run("PROJECT CANARY · TEAM STUDY GUIDE")
    set_font(kr, size=9, color=LIME, bold=True)

    title = doc.add_paragraph()
    paragraph_spacing(title, after=3, line=1.0)
    tr = title.add_run("Capstone Defense Cheat Sheet")
    set_font(tr, size=25, color=GREEN, bold=True)

    subtitle = doc.add_paragraph()
    paragraph_spacing(subtitle, after=10)
    sr = subtitle.add_run("What Canary is, how each component works, what the results mean, and which claims the evidence supports")
    set_font(sr, size=11.5, color=MUTED)

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    cell = callout.cell(0, 0)
    set_cell_fill(cell, LIGHT)
    p = cell.paragraphs[0]
    paragraph_spacing(p, after=0)
    add_inline(
        p,
        "**Core story:** Identify buildings that need attention, project Day 35 weight and harvest recovery, explain the warning signs, and recommend the next management check.",
        size=10.5,
    )
    spacer = doc.add_paragraph()
    paragraph_spacing(spacer, after=4)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = len(rows[0])
    width_map = {
        2: [3000, 6360],
        3: [2500, 1800, 5060],
        4: [2450, 1450, 2450, 3010],
    }
    widths = width_map.get(columns, [9360 // columns] * columns)
    widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for r_idx, source_row in enumerate(rows):
        row_props = table.rows[r_idx]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_props.append(cant_split)
        if r_idx == 0:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            row_props.append(header)
        for c_idx, value in enumerate(source_row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            paragraph_spacing(p, after=0, line=1.1)
            p.paragraph_format.keep_together = True
            if len(rows) <= 7 and r_idx < len(rows) - 1:
                p.paragraph_format.keep_with_next = True
            add_inline(p, value, size=9.2 if columns >= 3 else 9.6, color=WHITE if r_idx == 0 else INK)
            if r_idx == 0:
                set_cell_fill(cell, "286245")
                for run in p.runs:
                    run.bold = True
            elif r_idx % 2 == 0:
                set_cell_fill(cell, "F4F8F5")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build() -> Path:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    i = 2  # Markdown title and blank line are replaced by the designed title block.
    active_numbering_id: int | None = None
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            active_numbering_id = None
            i += 1
            continue
        if line.startswith("|"):
            table_rows: list[list[str]] = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    table_rows.append(cells)
                i += 1
            add_table(doc, table_rows)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("> "):
            table = doc.add_table(rows=1, cols=1)
            set_table_geometry(table, [9360])
            row_props = table.rows[0]._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            row_props.append(cant_split)
            cell = table.cell(0, 0)
            set_cell_fill(cell, "F2F8D9")
            p = cell.paragraphs[0]
            paragraph_spacing(p, after=0)
            add_inline(p, line[2:], size=10.5, color=GREEN)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif re.match(r"^\d+\. ", line):
            if active_numbering_id is None:
                active_numbering_id = new_numbering_sequence(doc)
            p = doc.add_paragraph()
            apply_numbering(p, active_numbering_id)
            add_inline(p, re.sub(r"^\d+\. ", "", line), size=10.5)
        elif line.startswith("- "):
            active_numbering_id = None
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, " " + line[2:], size=10.5)
        else:
            active_numbering_id = None
            p = doc.add_paragraph()
            paragraph_spacing(p, after=6, line=1.2)
            add_inline(p, line, size=10.5)
        i += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
