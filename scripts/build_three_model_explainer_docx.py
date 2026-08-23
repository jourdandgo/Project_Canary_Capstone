"""Build the formatted Project Canary three-model explainer DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Project_Canary_Three_Model_Explainer.docx"

# standard_business_brief preset, with a Project Canary palette override.
FONT = "Calibri"
GREEN = RGBColor(23, 63, 49)
GREEN_2 = RGBColor(40, 98, 69)
CANARY = RGBColor(150, 190, 43)
INK = RGBColor(38, 48, 43)
MUTED = RGBColor(92, 105, 99)
WHITE = RGBColor(255, 255, 255)
GOLD = RGBColor(122, 90, 0)
RED = RGBColor(155, 28, 28)

LIGHT_GREEN = "EDF5F0"
PALE_CANARY = "F6F9E8"
LIGHT_GRAY = "F2F4F7"
MID_GREEN = "286245"
DARK_GREEN = "173F31"
GRID = "D7E4DC"
WHITE_HEX = "FFFFFF"
MUTED_HEX = "5C6963"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    *,
    size: float = 11,
    color: RGBColor = INK,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = FONT
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph(
    paragraph,
    *,
    before: float = 0,
    after: float = 6,
    line: float = 1.10,
    keep_with_next: bool = False,
    keep_together: bool = False,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next
    fmt.keep_together = keep_together


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, *, color: str = GRID, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int], *, indent: int = TABLE_INDENT_DXA) -> None:
    assert sum(widths) == CONTENT_WIDTH_DXA
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr

    tbl_layout = table_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        table_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height = None
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run.append(instr)
    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    text_run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED_HEX)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "17")
    run_props.extend([color, size])
    text_run.append(run_props)
    text = OxmlElement("w:t")
    text.text = "1"
    text_run.append(text)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.extend([begin_run, instr_run, separate_run, text_run, end_run])


def add_numbering_definition(doc: Document, *, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, *, after: float = 4) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    set_paragraph(paragraph, after=after, line=1.10)


def add_text(paragraph, text: str, *, size: float = 11, color: RGBColor = INK, bold: bool = False, italic: bool = False) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_body(doc: Document, text: str, *, after: float = 6, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, after=after)
    if bold_lead and text.startswith(bold_lead):
        add_text(p, bold_lead, bold=True)
        add_text(p, text[len(bold_lead):])
    else:
        add_text(p, text)


def add_bullet(doc: Document, text: str, bullet_id: int, *, bold_lead: str | None = None, after: float = 4) -> None:
    p = doc.add_paragraph()
    apply_numbering(p, bullet_id, after=after)
    if bold_lead and text.startswith(bold_lead):
        add_text(p, bold_lead, bold=True)
        add_text(p, text[len(bold_lead):])
    else:
        add_text(p, text)


def add_numbered(doc: Document, text: str, number_id: int, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    apply_numbering(p, number_id, after=4)
    if bold_lead and text.startswith(bold_lead):
        add_text(p, bold_lead, bold=True)
        add_text(p, text[len(bold_lead):])
    else:
        add_text(p, text)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)


def add_kicker(doc: Document, text: str, *, new_page: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, after=3, keep_with_next=True)
    p.paragraph_format.page_break_before = new_page
    add_text(p, text.upper(), size=9, color=CANARY, bold=True)


def add_callout(doc: Document, label: str, text: str, *, fill: str = LIGHT_GREEN, accent: RGBColor = GREEN) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    keep_row_together(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph(p, after=1, line=1.10)
    add_text(p, f"{label}: ", size=10.5, color=accent, bold=True)
    add_text(p, text, size=10.5)
    spacer = doc.add_paragraph()
    set_paragraph(spacer, after=2)


def configure_section_geometry(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_header_footer(section) -> None:
    section.different_first_page_header_footer = False
    footer = section.footer
    add_page_field(footer.paragraphs[0])


def configure_document(doc: Document) -> tuple[int, int]:
    section = doc.sections[0]
    configure_section_geometry(section)
    section.different_first_page_header_footer = False

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, GREEN, 16, 8),
        "Heading 2": (13, GREEN_2, 12, 6),
        "Heading 3": (12, GREEN, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet_id = add_numbering_definition(doc, kind="bullet")
    number_id = add_numbering_definition(doc, kind="number")
    return bullet_id, number_id


def add_cover(doc: Document) -> None:
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph(p, after=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=14)
    add_text(p, "CAPSTONE MODEL EXPLAINER", size=9.5, color=CANARY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=8, line=1.0)
    add_text(p, "Project Canary", size=31, color=GREEN, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=6)
    add_text(p, "Three-Model Explainer", size=19, color=GREEN_2, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=30)
    add_text(p, "Two business outcomes. Three model engines. One traceable decision-support story.", size=12.5, color=MUTED, italic=True)

    add_callout(
        doc,
        "In one sentence",
        "Canary estimates the end-of-cycle recovery proxy and Day 35 bodyweight so farm management can see developing risk earlier, investigate sooner, and keep every forecast separate from the rules-based risk score.",
        fill=PALE_CANARY,
        accent=GREEN,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, before=22, after=3)
    add_text(p, "Prepared for JJ Agriventures", size=10.5, color=GREEN, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=0)
    add_text(p, "Pilot-stage model reference · August 2026", size=9.5, color=MUTED)
    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section_geometry(body_section)
    body_section.header.is_linked_to_previous = False
    body_section.footer.is_linked_to_previous = False
    set_header_footer(body_section)


def add_model_comparison_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    widths = [1920, 2360, 1900, 3180]
    set_table_geometry(table, widths)
    headers = ["Engine", "Outcome", "Forecast point", "Plain-language method"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, DARK_GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph(p, after=0, line=1.0)
        add_text(p, text, size=9.2, color=WHITE, bold=True)
    repeat_table_header(table.rows[0])
    rows = [
        ("Model 1\nExtra Trees", "End-of-cycle recovery proxy", "Day 7 and Day 14", "Averages forecasts from 500 decision trees using 85 engineered inputs."),
        ("Model 3\nXGBoost", "Day 35 bodyweight", "Day 21", "Uses 250 boosted trees and 88 engineered inputs available through Day 21."),
        ("Checkpoint model\nRemaining gain", "Day 35 bodyweight", "Days 7, 14, 21, and 28", "Adds the historical average remaining growth to the latest actual checkpoint weight."),
    ]
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        keep_row_together(table.rows[-1])
        for idx, value in enumerate(values):
            if row_idx % 2 == 0:
                shade_cell(cells[idx], LIGHT_GREEN)
            p = cells[idx].paragraphs[0]
            set_paragraph(p, after=0, line=1.05)
            for line_idx, line in enumerate(value.split("\n")):
                if line_idx:
                    p.add_run("\n")
                add_text(p, line, size=9.2, bold=(idx == 0 and line_idx == 0), color=GREEN if idx == 0 else INK)
    spacer = doc.add_paragraph()
    set_paragraph(spacer, after=3)


def add_metric_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    widths = [2160, 2460, 1680, 3060]
    set_table_geometry(table, widths)
    headers = ["Engine", "Development result", "2026-3 audit", "Current interpretation"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, DARK_GREEN)
        p = cell.paragraphs[0]
        set_paragraph(p, after=0, line=1.0)
        add_text(p, text, size=9, color=WHITE, bold=True)
    repeat_table_header(table.rows[0])
    data = [
        ("Model 1 · Extra Trees", "2.47 pp pooled MAE\n2.76 pp cycle-macro MAE", "4.55 pp MAE", "Experimental. The transparent recovery baseline was slightly better."),
        ("Model 3 · XGBoost", "146 g Day 21 MAE\n132 g cycle-macro MAE", "116 g MAE", "Useful Day 21 shadow benchmark, but later and less accurate at Day 21."),
        ("Checkpoint model", "127 g pooled MAE\n121 g cycle-macro MAE", "78 g MAE", "Provisional default for bodyweight; earlier, simpler, and stronger at Day 21."),
    ]
    for row_idx, values in enumerate(data):
        cells = table.add_row().cells
        keep_row_together(table.rows[-1])
        for idx, value in enumerate(values):
            if row_idx == 2:
                shade_cell(cells[idx], PALE_CANARY)
            elif row_idx % 2 == 0:
                shade_cell(cells[idx], LIGHT_GRAY)
            p = cells[idx].paragraphs[0]
            set_paragraph(p, after=0, line=1.05)
            lines = value.split("\n")
            for j, line in enumerate(lines):
                if j:
                    p.add_run("\n")
                color = GREEN if idx == 0 else INK
                add_text(p, line, size=9.0, color=color, bold=(idx == 0 or (row_idx == 2 and idx == 3)))


def add_checkpoint_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    widths = [2800, 2800, 3760]
    set_table_geometry(table, widths)
    headers = ["Forecast point", "Held-out MAE", "Lead time to Day 35"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, MID_GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(p, after=0)
        add_text(p, text, size=9.2, color=WHITE, bold=True)
    repeat_table_header(table.rows[0])
    for i, row in enumerate((("Day 7", "155 g", "28 days"), ("Day 14", "138 g", "21 days"), ("Day 21", "112 g", "14 days"), ("Day 28", "103 g", "7 days"))):
        cells = table.add_row().cells
        keep_row_together(table.rows[-1])
        for idx, value in enumerate(row):
            if i % 2 == 0:
                shade_cell(cells[idx], LIGHT_GREEN)
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph(p, after=0)
            add_text(p, value, size=9.5, bold=(idx == 0), color=GREEN if idx == 0 else INK)


def add_model_page(
    doc: Document,
    *,
    kicker: str,
    title: str,
    tagline: str,
    facts: list[tuple[str, str]],
    steps: list[tuple[str, str]],
    interpretation: str,
    status_label: str,
    status_text: str,
    number_id: int,
    status_fill: str,
    status_color: RGBColor,
    new_page: bool = True,
) -> None:
    add_kicker(doc, kicker, new_page=new_page)
    add_heading(doc, title, 1)
    add_body(doc, tagline, after=10)

    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [2200, 7160])
    for row_idx, (label, value) in enumerate(facts):
        cells = table.add_row().cells
        keep_row_together(table.rows[-1])
        shade_cell(cells[0], LIGHT_GREEN)
        if row_idx % 2 == 1:
            shade_cell(cells[1], LIGHT_GRAY)
        p = cells[0].paragraphs[0]
        set_paragraph(p, after=0)
        add_text(p, label, size=9.5, color=GREEN, bold=True)
        p = cells[1].paragraphs[0]
        set_paragraph(p, after=0)
        add_text(p, value, size=9.5)

    add_heading(doc, "How it was built", 2)
    for label, detail in steps:
        add_numbered(doc, f"{label} {detail}", number_id, bold_lead=f"{label} ")

    add_heading(doc, "How to interpret it", 2)
    add_body(doc, interpretation, after=8)
    add_callout(doc, status_label, status_text, fill=status_fill, accent=status_color)


def build_document() -> None:
    doc = Document()
    bullet_id, number_id = configure_document(doc)
    add_cover(doc)

    add_kicker(doc, "Executive overview")
    add_heading(doc, "What Project Canary is trying to do", 1)
    add_body(
        doc,
        "Project Canary gives farm management an earlier view of two possible end-of-cycle outcomes: the end-of-cycle recovery proxy and average bodyweight on Day 35. Earlier visibility creates more time to investigate a developing concern. It does not guarantee that production results will improve.",
    )
    add_callout(
        doc,
        "The key distinction",
        "There are two forecast outcomes, not three. Model 1 handles recovery. Model 3 and the checkpoint model are alternative engines for the same Day 35 bodyweight outcome.",
        fill=PALE_CANARY,
    )
    add_heading(doc, "The three model engines at a glance", 2)
    add_model_comparison_table(doc)
    add_heading(doc, "What the app should show", 2)
    for text, lead in [
        ("Recovery outlook: explicitly labelled “Model 1 · Extra Trees.”", "Recovery outlook:"),
        ("Bodyweight outlook: a toggle between “Checkpoint model” and “Model 3 · XGBoost.”", "Bodyweight outlook:"),
        ("Model details: evidence date, input freshness, forecast interval, validation error, method version, and pilot status.", "Model details:"),
        ("Risk-score separation: changing the forecast model must never change the observed-condition risk score.", "Risk-score separation:"),
    ]:
        add_bullet(doc, text, bullet_id, bold_lead=lead)
    add_callout(
        doc,
        "Trust boundary",
        "These are planning outlooks. They do not diagnose disease, prescribe treatment, trigger automatic action, or guarantee the final result.",
        fill=LIGHT_GRAY,
        accent=RED,
    )
    add_kicker(doc, "Shared foundation", new_page=True)
    add_heading(doc, "From raw farm records to a defensible forecast", 1)
    add_body(doc, "All three engines begin with the same cleaned building-day dataset and the same leakage-safe validation principle.")
    add_heading(doc, "Data preparation", 2)
    preparation = [
        ("Consolidate the records.", "Combine electronic and recovered physical records into one building-day dataset."),
        ("Standardize the fields.", "Harmonize names, dates, age, population, mortality, feed, bodyweight, temperature, humidity, units, and percentages."),
        ("Protect data meaning.", "Remove duplicate building-day rows and keep missing observations missing."),
        ("Define the outcomes.", "Use last recorded population divided by beginning population as the recovery proxy and the recorded average Day 35 bodyweight as the growth outcome."),
        ("Freeze the cohorts.", "Develop on 31 building-cycles from 2025-2 through 2026-2; reserve the three 2026-3 buildings for a later audit."),
    ]
    for label, detail in preparation:
        add_numbered(doc, f"{label} {detail}", number_id, bold_lead=f"{label} ")

    add_heading(doc, "Validation design", 2)
    add_body(doc, "One complete production cycle was held out at a time. Its buildings and daily observations stayed outside preprocessing, feature calculations, and fitting. This prevents related flock records from leaking into both training and testing.")
    add_callout(
        doc,
        "Why this matters",
        "A random row split would make performance look more confident than it should because daily observations from the same flock are closely related.",
        fill=LIGHT_GREEN,
    )
    add_heading(doc, "Feature engineering evaluated", 2)
    for text in [
        "Growth: weight gaps, ratios, interval gains, growth rates, trajectory, and projected Day 35 trajectory.",
        "Survival and mortality: daily loss, cumulative loss, current survival, and mortality patterns.",
        "Operations: feed use and gaps, population, stocking density, square feet per bird, downtime, farm, and building.",
        "Environment: temperature, humidity, ranges, deviations, and THI level, maximum, duration, and recent exposure.",
        "Data quality: missingness indicators and measurement freshness.",
    ]:
        add_bullet(doc, text, bullet_id)
    add_callout(
        doc,
        "Important caution",
        "Feature importance shows what a model relied on statistically. It does not prove that changing THI, feed, stocking density, or another input will cause the final outcome to improve.",
        fill=PALE_CANARY,
        accent=GOLD,
    )
    add_model_page(
        doc,
        kicker="Recovery forecast",
        title="Model 1 · Extra Trees",
        tagline="The reconstructed original approach for estimating the end-of-cycle recovery proxy.",
        facts=[
            ("Outcome", "End-of-cycle recovery proxy"),
            ("Forecast points", "Day 7 and Day 14; the Day 14 estimate is held afterward"),
            ("Algorithm", "Extra Trees regression with 500 trees"),
            ("Inputs", "85 locked engineered variables"),
            ("Development result", "2.47 percentage-point pooled MAE; 2.76-point cycle-macro MAE"),
            ("Later audit", "4.55 percentage-point MAE on the three buildings from 2026-3"),
        ],
        steps=[
            ("Prepare the evidence.", "Create a Day 7 and Day 14 snapshot for each eligible building-cycle without using later information."),
            ("Engineer the inputs.", "Describe survival, mortality, weight progress, feed, temperature, humidity, THI, housing context, missingness, and freshness."),
            ("Train the forest.", "Fit 500 randomized regression trees. Each tree learns different splits, and their predictions are averaged."),
            ("Validate by cycle.", "Hold out one complete production cycle at a time and calculate error only on unseen cycles."),
            ("Audit later data.", "Apply the frozen method once to the 2026-3 buildings."),
        ],
        interpretation="A higher forecast means the model expects a larger share of the beginning population to remain by the final recorded day. The value is still a proxy, not a verified harvest or sales count. Show the estimate with its uncertainty and evidence cutoff.",
        status_label="Current status",
        status_text="Experimental. The transparent recovery baseline had slightly lower held-out error, so Model 1 has not cleared the deployment gate. It is retained to evaluate the original approach, not to claim a proven champion.",
        number_id=add_numbering_definition(doc, kind="number"),
        status_fill="FDECEC",
        status_color=RED,
        new_page=False,
    )
    add_model_page(
        doc,
        kicker="Bodyweight forecast · shadow benchmark",
        title="Model 3 · XGBoost",
        tagline="A feature-rich model that estimates Day 35 bodyweight using information available through Day 21.",
        facts=[
            ("Outcome", "Recorded average bodyweight on Day 35"),
            ("Forecast point", "Day 21 only; the result is held afterward"),
            ("Algorithm", "XGBoost regression with 250 shallow boosted trees"),
            ("Inputs", "88 locked engineered variables through Day 21"),
            ("Development result", "146 g Day 21 MAE; 132 g cycle-macro MAE"),
            ("Later audit", "116 g MAE on the three buildings from 2026-3"),
        ],
        steps=[
            ("Set the evidence cutoff.", "Use only information available through Day 21."),
            ("Engineer the inputs.", "Describe growth trajectory, projected Day 35 trajectory, mortality, feed, temperature, humidity, THI, operations, missingness, and freshness."),
            ("Train sequentially.", "Fit 250 shallow trees. Each new tree focuses on prediction error left by the earlier trees."),
            ("Validate by cycle.", "Keep each held-out production cycle completely outside preprocessing and model fitting."),
            ("Audit later data.", "Run the frozen Day 21 method on the 2026-3 buildings."),
        ],
        interpretation="Model 3 offers a feature-rich Day 21 estimate with 14 days of lead time. Its broader input set may be useful for technical investigation, but those inputs should not be presented as proven causes of growth performance.",
        status_label="Current status",
        status_text="Shadow benchmark. It is available in the bodyweight toggle for comparison, but it should not be used before Day 21 or described as the default bodyweight forecast.",
        number_id=add_numbering_definition(doc, kind="number"),
        status_fill=LIGHT_GRAY,
        status_color=GREEN_2,
    )
    add_model_page(
        doc,
        kicker="Bodyweight forecast · provisional default",
        title="Checkpoint model · Historical remaining gain",
        tagline="A transparent method that updates after each actual weighing checkpoint and estimates the remaining growth to Day 35.",
        facts=[
            ("Outcome", "Recorded average bodyweight on Day 35"),
            ("Forecast points", "Days 7, 14, 21, and 28"),
            ("Calculation", "Latest actual weight + historical average remaining gain"),
            ("Direct inputs", "Checkpoint weight, measurement day, freshness, and training-fold remaining gain"),
            ("Development result", "127 g pooled MAE; 121 g cycle-macro MAE"),
            ("Later audit", "78 g MAE on the three buildings from 2026-3"),
        ],
        steps=[
            ("Read the latest measurement.", "Use an actual Day 7, 14, 21, or 28 bodyweight; do not invent a new weight between weigh-ins."),
            ("Estimate remaining gain.", "Calculate the historical average growth from that checkpoint to Day 35 using training cycles only."),
            ("Create the outlook.", "Add the estimated remaining gain to the latest actual checkpoint weight."),
            ("Validate by cycle.", "Recalculate the remaining-gain average without the held-out production cycle, then score the unseen buildings."),
            ("Compare candidates.", "Evaluate linear, regularized, robust, tree-based, boosting, CatBoost, XGBoost, and transparent baselines; retain the simplest stable winner."),
        ],
        interpretation="A heavier measured bird produces a higher Day 35 outlook. Later checkpoints are normally more accurate because less future growth remains unknown. Between weigh-ins, the forecast should remain unchanged while the app clearly shows the age of the latest measurement.",
        status_label="Current status",
        status_text="Provisional default for Day 35 bodyweight. It offers earlier visibility, refreshes four times, is easy to audit, and had lower Day 21 error than Model 3.",
        number_id=add_numbering_definition(doc, kind="number"),
        status_fill=PALE_CANARY,
        status_color=GREEN,
    )
    add_kicker(doc, "Evidence and operating choice", new_page=True)
    add_heading(doc, "How the three engines compare", 1)
    add_metric_table(doc)
    add_body(doc, "MAE is the average absolute difference between the forecast and the recorded outcome. Lower is better. Percentage-point error applies to recovery; gram error applies to bodyweight, so those values should not be compared with each other.", after=8)
    add_heading(doc, "Checkpoint accuracy improves as Day 35 approaches", 2)
    add_checkpoint_table(doc)
    add_body(doc, "At the same Day 21 evidence point, the checkpoint method's MAE was approximately 34 g lower than Model 3's: 112 g versus 146 g.", after=8, bold_lead="At the same Day 21 evidence point,")
    add_callout(
        doc,
        "Audit limitation",
        "The 2026-3 audit contains three buildings from one production cycle. It is a useful later-time check, but it is not three independent cycles and should not be treated as final proof of generalization.",
        fill=PALE_CANARY,
        accent=GOLD,
    )
    add_heading(doc, "When to use which", 2)
    for text, lead in [
        ("Model 1: show the original recovery-model approach at Day 7 or Day 14, clearly labelled experimental.", "Model 1:"),
        ("Checkpoint model: use as the default Day 35 bodyweight outlook because it is earlier, stronger at Day 21, and easier to audit.", "Checkpoint model:"),
        ("Model 3: retain as a Day 21 shadow benchmark when the team wants to compare the feature-rich XGBoost approach.", "Model 3:"),
        ("If bodyweight forecasts disagree: investigate the evidence and freshness; do not average them automatically and do not change the risk score.", "If bodyweight forecasts disagree:"),
    ]:
        add_bullet(doc, text, bullet_id, bold_lead=lead)
    add_callout(
        doc,
        "Recommended operating position",
        "Keep all three visible for capstone evaluation. Use Model 1 for recovery, make the checkpoint method the default bodyweight option, and keep Model 3 behind the bodyweight toggle as a clearly labelled Day 21 shadow benchmark.",
        fill=LIGHT_GREEN,
    )
    add_kicker(doc, "Presentation support")
    add_heading(doc, "Speaker notes", 1)
    add_body(doc, "The notes below are written to be read aloud or adapted into short presentation remarks.", after=10)
    speaker_notes = [
        ("Big picture", "Project Canary gives farm management an earlier view of where a flock may be heading. We forecast two outcomes: the end-of-cycle recovery proxy and average bodyweight on Day 35. We use three engines because two bodyweight approaches are being compared. Forecasts stay separate from the risk score and never trigger an automatic action."),
        ("Data and validation", "We consolidated the farm records into one building-day dataset, standardized names, dates, units, and percentages, removed duplicates, and preserved missing values. We developed the models on 31 building-cycles from six production cycles. For validation, we held out one complete production cycle at a time so related records could not leak into both training and testing. We opened 2026-3 only after the methods were frozen."),
        ("Model 1", "Model 1 predicts the end-of-cycle recovery proxy. It is an Extra Trees model with 500 decision trees and 85 engineered inputs. It produces outlooks at Day 7 and Day 14. Its historical error was about 2.47 percentage points, but it did not beat the transparent recovery baseline, so we label it experimental."),
        ("Model 3", "Model 3 predicts Day 35 bodyweight using information available through Day 21. It is an XGBoost model with 250 boosted trees and 88 engineered inputs, including growth trajectory, mortality, feed, temperature, humidity, and THI. Its Day 21 historical error was about 146 grams. It is a richer benchmark, but it provides only two weeks of lead time."),
        ("Checkpoint model", "The checkpoint model also predicts Day 35 bodyweight, but it works at Days 7, 14, 21, and 28. It takes the latest measured weight and adds the average historical growth that remained from that checkpoint to Day 35. Its overall error was about 127 grams, and at Day 21 its error was about 112 grams, compared with 146 grams for Model 3."),
        ("Why the toggle exists", "The toggle lets us compare the two bodyweight approaches without showing conflicting numbers at the same time. The selected model is written directly on the card. The checkpoint method is the default because it is earlier, more accurate at Day 21, and easier to audit. Model 3 remains a shadow benchmark."),
        ("THI and feature importance", "Temperature, humidity, and THI were retained in the engineered feature sets for Models 1 and 3. Feature importance tells us what the model relied on statistically; it does not prove that changing THI or another variable will cause the outcome to improve. We use those features as inspection context and research leads, not automatic treatment rules."),
        ("Trust boundary", "These models are suitable for a controlled shadow pilot and capstone comparison. They are not suitable for diagnosing disease, prescribing treatment, or promising production improvement. Model 1 remains experimental, Model 3 remains a shadow benchmark, and the checkpoint method is the provisional default bodyweight outlook."),
    ]
    for label, note in speaker_notes:
        add_callout(doc, label, note, fill=LIGHT_GRAY if label not in {"Big picture", "Trust boundary"} else LIGHT_GREEN, accent=GREEN)

    add_heading(doc, "One-line takeaway", 2)
    add_callout(
        doc,
        "Say this",
        "Canary uses one experimental recovery model and two bodyweight approaches to give management earlier, traceable planning outlooks—while keeping forecasts separate from the observed-condition risk score.",
        fill=PALE_CANARY,
        accent=GREEN,
    )
    add_body(doc, "Source: Project Canary three-model evaluation workflow and model artifacts. Metrics current as of August 2026.", after=0)

    # Core properties and deterministic behavior.
    doc.core_properties.title = "Project Canary: Three-Model Explainer"
    doc.core_properties.subject = "Business-friendly model explainer and speaker notes"
    doc.core_properties.author = "Project Canary Capstone Team"
    doc.core_properties.keywords = "Project Canary, recovery proxy, Day 35 bodyweight, Extra Trees, XGBoost"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
